# app.py  —  Single-file Streamlit app (Windows-friendly, $0 runtime)
# Prereqs: pip install streamlit pydantic sentence-transformers chromadb pypdf python-docx jinja2 loguru requests
# Start:  streamlit run app.py

import json, re, hashlib
from pathlib import Path
from typing import List, Dict, Set, Tuple
from io import BytesIO

import requests
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------- Config ----------
DATA_DIR = Path("data")
OUT_DIR = Path("out")
CHROMA_DIR = DATA_DIR / "chroma_store"
OLLAMA_URL = "http://localhost:11434/api/chat"
OLLAMA_MODEL = "llama3.2:3b"  # Faster model; use "llama3.1:8b" if you need more capability
EMB_MODEL_NAME = "sentence-transformers/all-MiniLM-L6-v2"

# ---------- Utilities ----------
def _hash(s: str) -> str:
    return hashlib.sha256(s.encode("utf-8")).hexdigest()

def _clean_line(s: str) -> str:
    s = s.strip()
    s = re.sub(r"[;,:.\-—]+$", "", s).strip()
    return s

def _sentence_case(s: str) -> str:
    s = s.strip()
    return s[:1].upper() + s[1:] if s else s

# ---------- Skills Normalization (inline version of normalize_skills.py) ----------
def load_skills_map(path: str | Path = DATA_DIR / "skills_map.json") -> Dict:
    path = Path(path)
    if not path.exists():
        return {}
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)

def normalize_text(text: str, skills_map: Dict) -> tuple[str, Set[str]]:
    detected = set()
    normalized_text = text.lower()
    for canonical, synonyms in skills_map.items():
        for syn in synonyms:
            pattern = r"\b" + re.escape(syn.lower()) + r"\b"
            if re.search(pattern, normalized_text):
                normalized_text = re.sub(pattern, canonical.lower(), normalized_text)
                detected.add(canonical.lower())
        if re.search(r"\b" + re.escape(canonical.lower()) + r"\b", normalized_text):
            detected.add(canonical.lower())
    return normalized_text, detected

# ---------- JD Parser (inline, uses normalizer) ----------
_STOPWORDS = {
    "the","and","or","to","of","in","on","for","with","a","an","by","as","is","are","be",
    "will","this","that","we","you","our","your","at","from","across","including","such",
    "etc","about","into","within","over","under","more","most","least","ability","experience",
    "years","year","plus","etc.","responsibilities","requirements","preferred","must","nice",
    "have","skills","role","job","description","position","candidate","ideal","work","working"
}
_SECTION_HEADERS = {
    "responsibilities": r"(responsibilities|what you'll do|what you will do|duties|role and responsibilities)",
    "must_haves": r"(requirements|must[-\s]*have|qualifications|required)",
    "nice_to_haves": r"(nice[-\s]*to[-\s]*have|preferred|bonus|good to have)"
}
_SENIORITY_MAP = [
    (r"\b(principal|lead|staff)\b", "Principal/Lead"),
    (r"\b(senior|sr\.)\b", "Senior"),
    (r"\b(mid[-\s]*level|intermediate)\b", "Mid"),
    (r"\b(junior|jr\.)\b", "Junior"),
]

def _extract_title(text: str) -> str:
    m = re.search(r"(?im)^\s*title\s*[:\-]\s*(.+)$", text)
    if m: return _clean_line(m.group(1))
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    if lines:
        first = lines[0]
        if len(first.split()) <= 12 and not re.search(r"(company|about us|who we are)", first, re.I):
            return _clean_line(first)
    return ""

def _extract_company(text: str) -> str:
    m = re.search(r"(?im)^\s*company\s*[:\-]\s*(.+)$", text)
    if m: return _clean_line(m.group(1))
    m = re.search(r"(?im)about\s+([A-Z][A-Za-z0-9&\-\s]{2,40})", text)
    return _clean_line(m.group(1)) if m else ""

def _extract_location(text: str) -> str:
    m = re.search(r"(?im)^\s*location\s*[:\-]\s*(.+)$", text)
    if m: return _clean_line(m.group(1))
    m = re.search(r"(?i)\b(singapore|london|new york|sydney|remote|hybrid|onsite)\b", text)
    return m.group(0).title() if m else ""

def _guess_seniority(text: str) -> str:
    for pat, label in _SENIORITY_MAP:
        if re.search(pat, text, re.I):
            return label
    return ""

def _extract_bullets(text: str, header_regex: str) -> List[str]:
    header = re.search(header_regex, text, re.I)
    if not header: return []
    section = text[header.start():]
    m = re.search(r"(?s)(?:\r?\n|\r)\s*([\-*\u2022].+?)(?:\n\s*\n|$)", section)
    if not m: return []
    bullets_block = m.group(1)
    lines = [l.strip() for l in bullets_block.splitlines()]
    out = []
    for line in lines:
        if re.match(r"^\s*[\-*\u2022]\s+", line):
            cleaned = re.sub(r"^\s*[\-*\u2022]\s+", "", line).strip()
            if cleaned: out.append(_clean_line(cleaned))
    return out

def _normalize_lines(lines: List[str], skills_map: Dict) -> List[str]:
    out = []
    for l in lines:
        normalized, _ = normalize_text(l, skills_map)
        out.append(_sentence_case(normalized))
    return out

def _top_keywords(text: str, n: int = 30) -> List[str]:
    words = re.findall(r"[A-Za-z0-9\+\#\.]{3,}", text.lower())
    freq = {}
    for w in words:
        if w in _STOPWORDS: continue
        freq[w] = freq.get(w, 0) + 1
    return [w for w, _ in sorted(freq.items(), key=lambda x: -x[1])[:n]]

def parse_jd(jd_text: str, skills_map: Dict, top_kw: int = 30) -> Dict:
    jd_text = jd_text.strip()
    title = _extract_title(jd_text)
    company = _extract_company(jd_text)
    location = _extract_location(jd_text)
    seniority = _guess_seniority(jd_text)
    must_haves_raw = _extract_bullets(jd_text, _SECTION_HEADERS["must_haves"])
    nice_to_haves_raw = _extract_bullets(jd_text, _SECTION_HEADERS["nice_to_haves"])
    responsibilities_raw = _extract_bullets(jd_text, _SECTION_HEADERS["responsibilities"])
    normalized_text, detected_skills = normalize_text(jd_text, skills_map)
    must_haves = _normalize_lines(must_haves_raw, skills_map)
    nice_to_haves = _normalize_lines(nice_to_haves_raw, skills_map)
    responsibilities = _normalize_lines(responsibilities_raw, skills_map)
    keywords = _top_keywords(normalized_text, n=top_kw)
    return {
        "title": title,
        "company": company,
        "location": location,
        "seniority": seniority,
        "must_haves": must_haves[:15],
        "nice_to_haves": nice_to_haves[:15],
        "responsibilities": responsibilities[:30],
        "keywords": keywords,
        "normalized_skills": sorted(detected_skills),
        "normalized_text": normalized_text
    }

# ---------- Retrieval with Chroma + skill-overlap boost ----------
def ensure_chroma_and_index() -> "chromadb.api.models.Collection.Collection":
    import chromadb
    from sentence_transformers import SentenceTransformer

    client = chromadb.PersistentClient(path=str(CHROMA_DIR))
    names = [c.name for c in client.list_collections()]
    if "snippets" not in names:
        client.create_collection("snippets")
    coll = client.get_collection("snippets")

    if coll.count() == 0:
        items, ids, metas = [], [], []
        snip_path = DATA_DIR / "snippets.jsonl"
        if not snip_path.exists():
            return coll
        with snip_path.open("r", encoding="utf-8") as f:
            for line in f:
                obj = json.loads(line)
                items.append(obj["text"])
                ids.append(obj["id"])
                metas.append({"tags": obj.get("tags", []), "seniority": obj.get("seniority", "")})
        coll.add(ids=ids, documents=items, metadatas=metas)
    return coll

def _overlap_score(snippet_meta: Dict, jd_skills: Set[str]) -> float:
    tags = set((snippet_meta or {}).get("tags", []))
    tags = {t.lower() for t in tags}
    inter = len(tags & jd_skills)
    union = len(tags | jd_skills) or 1
    return inter / union

def query_snippets(jd_text_or_norm: str, jd_skills: Set[str] | None = None, top_k: int = 10) -> List[str]:
    coll = ensure_chroma_and_index()
    # query more and re-rank
    res = coll.query(query_texts=[jd_text_or_norm], n_results=min(50, top_k * 2),
                     include=["documents", "metadatas", "distances"])
    docs = res.get("documents", [[]])[0]
    metas = res.get("metadatas", [[]])[0]
    dists = res.get("distances", [[]])[0] or [1.0] * len(docs)

    base_scores = [1.0 / (1.0 + d) for d in dists]
    jd_skills = {s.lower() for s in (jd_skills or set())}

    boosted = []
    for doc, meta, base in zip(docs, metas, base_scores):
        bonus = _overlap_score(meta, jd_skills)
        boosted.append((doc, meta, base + 0.25 * bonus))
    boosted.sort(key=lambda x: x[2], reverse=True)
    return [b[0] for b in boosted[:top_k]]

# ---------- LLM (Ollama) with streaming ----------
SYS = (
"You are a meticulous career documents assistant. "
"Only use facts from MASTER_PROFILE and RETRIEVED_SNIPPETS. "
"Do NOT invent employers, dates, tools, or metrics. "
"Prefer active voice; bullets ≤ 22 words. "
"Output EXACTLY in the requested format."
)

def ollama_chat(user_content: str) -> str:
    """Call Ollama with streaming to avoid timeouts and show progress."""
    try:
        r = requests.post(OLLAMA_URL, json={
            "model": OLLAMA_MODEL,
            "messages": [{"role": "system", "content": SYS},
                         {"role": "user", "content": user_content}],
            "stream": True  # Enable streaming
        }, timeout=600, stream=True)  # 10 minute timeout
        r.raise_for_status()
        
        full_response = ""
        for line in r.iter_lines():
            if line:
                try:
                    json_response = json.loads(line)
                    if "message" in json_response:
                        content = json_response["message"].get("content", "")
                        full_response += content
                    if json_response.get("done", False):
                        break
                except json.JSONDecodeError:
                    continue
        
        return full_response
    
    except requests.exceptions.Timeout:
        st.error("⏱️ Ollama request timed out. Try: 1) Using a smaller model (llama3.2:3b), 2) Reducing snippet count, or 3) Ensuring Ollama is running properly.")
        raise
    except requests.exceptions.ConnectionError:
        st.error("🔌 Cannot connect to Ollama. Make sure Ollama is running: `ollama serve`")
        raise
    except Exception as e:
        st.error(f"❌ Ollama error: {str(e)}")
        raise

# ---------- Exporters ----------
def write_resume_docx(cv: Dict, path: Path = None) -> BytesIO:
    """Create resume DOCX in memory or save to path based on template selection."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Check if this is template-based generation
    if "template" in cv and "selection" in cv:
        template = cv["template"]
        selection = cv["selection"]
        
        # Header
        header = template.get("header", {})
        if header:
            name_para = doc.add_heading(header.get("name", ""), level=0)
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            title_para = doc.add_paragraph(header.get("title_line", ""))
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            contact_info = f"E: {header.get('email', '')} | T: {header.get('phone', '')} | IN: {header.get('linkedin', '')}"
            contact_para = doc.add_paragraph(contact_info)
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # spacing
        
        # Summary
        summary_section = template.get("summary", {})
        if summary_section:
            doc.add_heading("SUMMARY", level=1)
            doc.add_paragraph(summary_section.get("paragraph", ""))
            
            # Add selected summary bullets
            bullets = summary_section.get("bullets", [])
            bullets_to_keep = selection.get("summary_bullets_to_keep", [])
            for idx in bullets_to_keep:
                if 0 <= idx < len(bullets):
                    p = doc.add_paragraph(bullets[idx])
                    p.style = doc.styles["List Bullet"]
            
            doc.add_paragraph()  # spacing
        
        # Key Skills
        skills_section = template.get("key_skills", {})
        skills_to_highlight = selection.get("skills_to_highlight", [])
        
        if skills_section:
            doc.add_heading("KEY SKILLS", level=1)
            
            # Create a table for skills (matches your CV format)
            if skills_to_highlight:
                # Show only highlighted skills
                for skill_key in skills_to_highlight:
                    if skill_key in skills_section:
                        p = doc.add_paragraph()
                        p.add_run(f"{skill_key}").bold = True
                        skills_list = ", ".join(skills_section[skill_key])
                        doc.add_paragraph(skills_list)
            else:
                # Show all skills
                for skill_category, skills_list in skills_section.items():
                    p = doc.add_paragraph()
                    p.add_run(f"{skill_category}").bold = True
                    doc.add_paragraph(", ".join(skills_list))
            
            doc.add_paragraph()  # spacing
        
        # Experience
        experience_list = template.get("experience", [])
        experience_selection = selection.get("experience", [])
        
        if experience_list:
            doc.add_heading("EXPERIENCE", level=1)
            
            for exp_sel in experience_selection:
                exp_id = exp_sel.get("id")
                include = exp_sel.get("include", True)
                
                if not include:
                    continue
                
                # Find matching experience in template
                exp = next((e for e in experience_list if e.get("id") == exp_id), None)
                if not exp:
                    continue
                
                # Company header
                header_text = f"{exp.get('company', '')} | {exp.get('title', '')} | {exp.get('dates', '')}"
                doc.add_heading(header_text, level=2)
                
                # Add selected bullets
                bullets = exp.get("bullets", [])
                bullets_to_keep = exp_sel.get("bullets_to_keep", [])
                for idx in bullets_to_keep:
                    if 0 <= idx < len(bullets):
                        p = doc.add_paragraph(bullets[idx])
                        p.style = doc.styles["List Bullet"]
                
                # Add achievements section if any selected
                achievements = exp.get("achievements", [])
                achievements_to_keep = exp_sel.get("achievements_to_keep", [])
                
                if achievements_to_keep:
                    doc.add_heading("Achievements", level=3)
                    for idx in achievements_to_keep:
                        if 0 <= idx < len(achievements):
                            p = doc.add_paragraph(achievements[idx])
                            p.style = doc.styles["List Bullet"]
                
                doc.add_paragraph()  # spacing
        
        # Training and Education
        training_list = template.get("training_and_education", [])
        training_to_keep = selection.get("training_and_education_to_keep", [])
        
        if training_list:
            doc.add_heading("TRAINING AND EDUCATION", level=1)
            
            if training_to_keep:
                for idx in training_to_keep:
                    if 0 <= idx < len(training_list):
                        p = doc.add_paragraph(training_list[idx])
                        p.style = doc.styles["List Bullet"]
            else:
                # Show all if none selected
                for item in training_list:
                    p = doc.add_paragraph(item)
                    p.style = doc.styles["List Bullet"]
    
    else:
        # Legacy format (backward compatibility)
        if cv.get("headline"):
            doc.add_heading(cv["headline"], level=0)
        
        if cv.get("summary"):
            doc.add_paragraph(cv["summary"])

        if cv.get("skills"):
            doc.add_heading("Skills", level=1)
            doc.add_paragraph(", ".join(cv["skills"]))

        if cv.get("experience"):
            doc.add_heading("Experience", level=1)
            for exp in cv["experience"]:
                header = f'{exp.get("role","")} — {exp.get("company","")} ({exp.get("dates","")})'
                doc.add_paragraph(header)
                for b in exp.get("bullets", []):
                    p = doc.add_paragraph(b)
                    p.style = doc.styles["List Bullet"]

        if cv.get("education"):
            doc.add_heading("Education", level=1)
            for e in cv["education"]:
                doc.add_paragraph(e)

        if cv.get("certifications"):
            doc.add_heading("Certifications", level=1)
            for c in cv["certifications"]:
                doc.add_paragraph(c)

    # Save to BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Also save to path if provided
    if path:
        doc.save(str(path))
    
    return buffer

# ---------- Prompts ----------
def build_cv_prompt(mp: Dict, jd: Dict, snips: List[str], template_profile: Dict) -> str:
    """
    Ask the model to ONLY SELECT which parts of the template CV to keep,
    instead of rewriting the whole CV from scratch.
    """
    jd_for_prompt = {
        "title": jd.get("title", ""),
        "company": jd.get("company", ""),
        "location": jd.get("location", ""),
        "seniority": jd.get("seniority", ""),
        "must_haves": jd.get("must_haves", []),
        "nice_to_haves": jd.get("nice_to_haves", []),
        "normalized_skills": jd.get("normalized_skills", []),
        "normalized_text": jd.get("normalized_text", "")
    }

    return f"""
You are a CV tailoring assistant.

You are given:
1) TEMPLATE_CV: a complete CV for one person, in JSON, that already has the correct wording, sections and layout.
2) MASTER_PROFILE: structured data about the same person (used only to understand context).
3) JOB_DESCRIPTION: the target job.
4) SNIPPETS: extra tagged bullets derived from the same CV.

Your job is ONLY to SELECT which parts of TEMPLATE_CV to keep for this application.

IMPORTANT RULES:
- DO NOT invent new employers, dates, tools, or metrics.
- DO NOT create new bullets from scratch.
- DO NOT change section titles or their order.
- DO NOT significantly rewrite text. You may only remove bullets or keep them.
- If a bullet is slightly misaligned with the job, prefer DROPPING it rather than rewriting it.
- Prefer the most senior and most recent experience when choosing what to keep.

TEMPLATE_CV:
{json.dumps(template_profile, ensure_ascii=False, indent=2)}

MASTER_PROFILE:
{json.dumps(mp, ensure_ascii=False, indent=2)}

JOB_DESCRIPTION:
{json.dumps(jd_for_prompt, ensure_ascii=False, indent=2)}

SNIPPETS (top {len(snips)}):
{json.dumps(snips, ensure_ascii=False, indent=2)}

Output ONLY a JSON object in this exact schema:

{{
  "summary_bullets_to_keep": [int, int, ...],

  "experience": [
    {{
      "id": "insead",
      "include": true or false,
      "bullets_to_keep": [int, int, ...],
      "achievements_to_keep": [int, int, ...]
    }},
    {{
      "id": "moqdigital",
      "include": true or false,
      "bullets_to_keep": [int, int, ...],
      "achievements_to_keep": [int, int, ...]
    }},
    {{
      "id": "pearson",
      "include": true or false,
      "bullets_to_keep": [int, int, ...],
      "achievements_to_keep": [int, int, ...]
    }},
    {{
      "id": "99x",
      "include": true or false,
      "bullets_to_keep": [int, int, ...],
      "achievements_to_keep": [int, int, ...]
    }},
    {{
      "id": "lanka_comm",
      "include": true or false,
      "bullets_to_keep": [int, int, ...],
      "achievements_to_keep": [int, int, ...]
    }},
    {{
      "id": "eyepax",
      "include": true or false,
      "bullets_to_keep": [int, int, ...],
      "achievements_to_keep": [int, int, ...]
    }}
  ],

  "skills_to_highlight": [
    "exact KEYS from TEMPLATE_CV.key_skills that are most relevant to the job",
    "e.g. 'Scripting & Automation', 'Cloud Technologies'"
  ],

  "training_and_education_to_keep": [int, int, ...]
}}
"""


def build_cl_prompt(mp: Dict, jd: Dict, snips: List[str]) -> str:
    return f"""
Task: Write a concise cover letter tailored to the JD.
Structure: Hook, Alignment (2–3 themes), Evidence (2–3 quantified outcomes), Close.
Tone: professional, confident, succinct.
Length: 180–300 words.
Include the role and company name.

MASTER_PROFILE:
{json.dumps(mp, ensure_ascii=False)}

JD (normalized):
Title: {jd.get("title","")}
Company: {jd.get("company","")}
Location: {jd.get("location","")}
Seniority: {jd.get("seniority","")}
Must-haves: {jd.get("must_haves", [])}
Nice-to-haves: {jd.get("nice_to_haves", [])}
Skills (normalized): {jd.get("normalized_skills", [])}
Text: {jd.get("normalized_text","")}

RETRIEVED_SNIPPETS (top {len(snips)}):
{json.dumps(snips, ensure_ascii=False)}
"""

if "cv" not in st.session_state:
    st.session_state["cv"] = None
if "letter" not in st.session_state:
    st.session_state["letter"] = None
if "jd" not in st.session_state:
    st.session_state["jd"] = None
if "snips" not in st.session_state:
    st.session_state["snips"] = None
if "jd_hash" not in st.session_state:
    st.session_state["jd_hash"] = None

# ---------- Streamlit UI ----------
st.set_page_config(page_title="CV Tailor (Local, Windows)", layout="wide")
st.title("🚀 Local CV & Cover Letter Tailor (Windows, $0 runtime)")

# Sidebar: master profile
st.sidebar.header("⚙️ Setup")
MP_FILE = DATA_DIR / "master_profile.json"
TEMPLATE_FILE = DATA_DIR / "template_profile.json"

skills_map = load_skills_map(DATA_DIR / "skills_map.json")
if not skills_map:
    st.sidebar.warning("⚠️ skills_map.json not found. Normalization will be skipped.")

mp = json.load(MP_FILE.open("r", encoding="utf-8")) if MP_FILE.exists() else None
st.sidebar.write("**Master Profile:**", "✅ Loaded" if mp else "❌ Missing")

# Load the fixed CV template (your current CV structure)
template_profile = json.load(TEMPLATE_FILE.open("r", encoding="utf-8")) if TEMPLATE_FILE.exists() else None
st.sidebar.write("**Template CV:**", "✅ Loaded" if template_profile else "❌ Missing")


uploaded = st.sidebar.file_uploader("Upload/Replace Master Profile (JSON)", type="json")
if uploaded:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    MP_FILE.write_bytes(uploaded.read())
    st.sidebar.success("✅ Saved. Reload the page to use the new profile.")

st.sidebar.markdown("---")
st.sidebar.markdown(f"**Model:** `{OLLAMA_MODEL}`")
st.sidebar.markdown(f"**Ollama URL:** `{OLLAMA_URL}`")

# Main input
st.markdown("### 📋 Job Description")
jd_text = st.text_area("Paste the full Job Description here:", height=320, 
                       placeholder="Paste the JD here...\n\nTitle: Senior DevOps Engineer\nCompany: TechCorp\n...")

col1, col2, col3 = st.columns([2, 1, 1])
with col1:
    gen_clicked = st.button("🎯 Generate CV & Cover Letter", type="primary", use_container_width=True)
with col2:
    show_debug = st.checkbox("🔍 Debug info", value=False)
with col3:
    snippet_count = st.number_input("Snippets", min_value=5, max_value=30, value=10, step=5)

if gen_clicked:
    if not mp:
        st.error("❌ Upload master_profile.json first (sidebar).")
        st.stop()
    if not jd_text.strip():
        st.error("❌ Paste a Job Description first.")
        st.stop()
    if not template_profile:
        st.error("❌ template_profile.json not found in data/. Create it first from your current CV.")
        st.stop()

    # ---------- Parse JD & retrieve snippets ----------
    with st.spinner("🔄 Parsing JD and retrieving snippets..."):
        # Hash & (re)parse/retrieve on change
        jd_hash = _hash(jd_text)
        if st.session_state.get("jd_hash") != jd_hash:
            jd_parsed = parse_jd(jd_text, skills_map)
            norm_text = jd_parsed.get("normalized_text") or jd_text
            norm_skills = set(jd_parsed.get("normalized_skills", []))
            snips = query_snippets(norm_text, norm_skills, top_k=snippet_count)

            st.session_state["jd"] = jd_parsed
            st.session_state["snips"] = snips
            st.session_state["jd_hash"] = jd_hash

        jd = st.session_state["jd"]
        snips = st.session_state["snips"]

    if show_debug:
        with st.expander("🔍 Debug Info"):
            st.caption(f"**JD Hash:** {jd_hash[:12]}...")
            st.write(f"**Title:** {jd.get('title', 'N/A')}")
            st.write(f"**Company:** {jd.get('company', 'N/A')}")
            st.write(f"**Location:** {jd.get('location', 'N/A')}")
            st.write(f"**Seniority:** {jd.get('seniority', 'N/A')}")
            st.write(f"**Detected skills ({len(jd.get('normalized_skills', []))}):**", 
                    ", ".join(jd.get("normalized_skills", [])[:20]))
            st.write("**Top 5 snippets:**")
            for i, s in enumerate(snips[:5], 1):
                st.text(f"{i}. {s[:100]}...")

    # ---------- Build prompts and call LLM for SELECTION ----------
    st.markdown("---")
    st.markdown("### 📄 Generating Resume (Template Selection)...")
    
    with st.spinner(f"🤖 Calling Ollama ({OLLAMA_MODEL})... This may take 1–3 minutes."):
        cv_prompt = build_cv_prompt(mp, jd, snips, template_profile)
        
        if show_debug:
            with st.expander("🔍 CV Prompt Preview"):
                st.text_area("Full prompt sent to LLM:", cv_prompt, height=200)
        
        try:
            selection_raw = ollama_chat(cv_prompt).strip()
        except Exception as e:
            st.error(f"Failed to generate CV selection: {e}")
            st.stop()

    # ---------- Parse SELECTION JSON (indices etc.) ----------
    selection_str = re.sub(r"^```json\s*|```\s*$", "", selection_raw, flags=re.M).strip()
    try:
        selection = json.loads(selection_str)
    except Exception as parse_err:
        # last resort: try to extract the largest {...} block
        m = re.search(r"\{[\s\S]+\}", selection_raw)
        if m:
            try:
                selection = json.loads(m.group(0))
            except Exception:
                selection = {
                    "summary_bullets_to_keep": [],
                    "experience": [],
                    "skills_to_highlight": [],
                    "training_and_education_to_keep": []
                }
        else:
            selection = {
                "summary_bullets_to_keep": [],
                "experience": [],
                "skills_to_highlight": [],
                "training_and_education_to_keep": []
            }
            st.warning("⚠️ Could not parse LLM response as JSON. Showing raw output:")
            st.code(selection_raw)

    # Wrap into cv structure that write_resume_docx() understands in template-mode
    cv = {
        "template": template_profile,
        "selection": selection
    }

    st.success("✅ Resume selection generated!")
    with st.expander("📋 Selection JSON", expanded=True):
        st.json(selection)

    # ---------- Generate Cover Letter (unchanged) ----------
    st.markdown("### 💌 Generating Cover Letter...")
    with st.spinner("🤖 Writing cover letter..."):
        cl_prompt = build_cl_prompt(mp, jd, snips)
        
        if show_debug:
            with st.expander("🔍 Cover Letter Prompt Preview"):
                st.text_area("Full prompt sent to LLM:", cl_prompt, height=200)
        
        try:
            letter = ollama_chat(cl_prompt).strip()
        except Exception as e:
            st.error(f"Failed to generate cover letter: {e}")
            letter = "Cover letter generation failed."

        st.session_state["letter"] = letter
        st.session_state["cv"] = cv

    st.success("✅ Cover letter generated!")
    with st.expander("💌 Cover Letter", expanded=True):
        st.markdown(letter)

    # ---------- Quality checks ----------
    st.markdown("---")
    st.markdown("### ⚠️ Quality Checks")
    missing_must_haves = [
        k for k in jd.get("must_haves", []) 
        if k and (k.lower() not in json.dumps(cv).lower() and k.lower() not in letter.lower())
    ]
    if missing_must_haves:
        st.warning(f"⚠️ **Missing must-have keywords:** {', '.join(missing_must_haves[:5])}")
    else:
        st.success("✅ All must-have keywords covered!")

    # ---------- Exports ----------
    st.markdown("---")
    st.markdown("### 💾 Download Documents")
    
    OUT_DIR.mkdir(exist_ok=True)
    
    col_dl1, col_dl2 = st.columns(2)
    
    with col_dl1:
        try:
            docx_buffer = write_resume_docx(cv, OUT_DIR / "resume_tailored.docx")
            
            st.download_button(
                label="📄 Download Resume (DOCX)",
                data=docx_buffer.getvalue(),
                file_name="resume_tailored.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        except Exception as e:
            st.error(f"Error creating resume DOCX: {e}")
            if show_debug:
                import traceback
                st.code(traceback.format_exc())
    
    with col_dl2:
        st.download_button(
            label="📝 Download Cover Letter (TXT)",
            data=letter.encode("utf-8"),
            file_name="cover_letter.txt",
            mime="text/plain",
            use_container_width=True
        )
    
    st.success("✅ Documents ready for download!")
    st.info("💡 **Tip:** Review the documents before submitting. Personalize the cover letter opening/closing if needed.")
