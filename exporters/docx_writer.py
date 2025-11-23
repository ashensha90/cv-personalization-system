# exporters/docx_writer.py
from docx import Document
from docx.shared import Pt

def write_resume_docx(cv, path):
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(11)

    doc.add_heading(cv["headline"], level=0)
    doc.add_paragraph(cv["summary"])

    doc.add_heading("Skills", level=1)
    doc.add_paragraph(", ".join(cv["skills"]))

    doc.add_heading("Experience", level=1)
    for exp in cv["experience"]:
        doc.add_paragraph(f'{exp["role"]} — {exp["company"]} ({exp["dates"]})')
        for b in exp["bullets"]:
            doc.add_paragraph(b, style=None).style = doc.styles["List Bullet"]

    if cv.get("education"):
        doc.add_heading("Education", level=1)
        for e in cv["education"]:
            doc.add_paragraph(e)

    if cv.get("certifications"):
        doc.add_heading("Certifications", level=1)
        for c in cv["certifications"]:
            doc.add_paragraph(c)

    doc.save(path)
